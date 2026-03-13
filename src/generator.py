from docx import Document
from docx.shared import Inches
import os

class DDRGenerator:
    def __init__(self, output_path="output/Final_DDR_Report.docx"):
        self.output_path = output_path
        self.doc = Document()

    def create_report(self, report_text, image_mappings=None):
        self.doc.add_heading('Detailed Diagnostic Report (DDR)', 0)

        # Split report into sections based on the 7 required points
        sections = report_text.split('###')
        
        for section in sections:
            section_content = section.strip()
            if not section_content:
                continue
                
            self.doc.add_paragraph(section_content)

            # Better Image Placement Logic
            if image_mappings:
                for mapping in image_mappings:
                    # Check if 'area' key exists; if not, use page number as a reference
                    area_name = mapping.get('area', f"Page {mapping.get('page', 'Unknown')}")
                    
                    # If the section mentions the area OR if we are in 'Area-wise Observations'
                    if area_name.lower() in section_content.lower() or "area-wise" in section_content.lower():
                        if os.path.exists(mapping['path']):
                            self.doc.add_picture(mapping['path'], width=Inches(3))
                            self.doc.add_paragraph(f"Figure: Observation from {area_name}")
                        else:
                            self.doc.add_paragraph("Image Not Available") # Requirement [cite: 58]
        
        self.doc.save(self.output_path)
        print(f"✅ Report saved to {self.output_path}")